"""Extração e chunking de artefatos anexados a conversas."""

from __future__ import annotations

import io
import re
import uuid
from dataclasses import dataclass
from pathlib import PurePath

from app.domain.entities import ConversationArtifactChunk

IMAGE_MIME_TYPES = {
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/webp",
    "image/gif",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".log",
    ".json",
    ".yaml",
    ".yml",
    ".csv",
    ".xml",
}

TEXT_MIME_PREFIXES = ("text/",)
TEXT_MIME_TYPES = {
    "application/json",
    "application/xml",
    "application/x-yaml",
    "application/yaml",
}

TEXT_CHUNK_CHARS = 4000
TEXT_CHUNK_OVERLAP = 300
LOG_CHUNK_LINES = 200
LOG_CHUNK_CHARS = 12000


class ArtifactProcessingError(ValueError):
    """Falha de extração ou tipo de arquivo não suportado."""


@dataclass(frozen=True)
class ArtifactProcessingResult:
    kind: str
    chunks: list[ConversationArtifactChunk]


def detect_artifact_kind(filename: str, mime_type: str) -> str:
    """Classifica o artefato a partir de MIME e extensão."""
    mime = mime_type.lower().split(";")[0].strip()
    suffix = PurePath(filename.lower()).suffix
    if mime in IMAGE_MIME_TYPES or suffix in IMAGE_EXTENSIONS:
        return "image"
    if suffix == ".pdf" or mime == "application/pdf":
        return "pdf"
    if suffix == ".docx" or mime.endswith(
        "vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return "docx"
    if suffix == ".doc" or mime == "application/msword":
        raise ArtifactProcessingError(
            "Arquivos .doc legados não são suportados. Converta para .docx ou PDF."
        )
    if suffix == ".log":
        return "log"
    if suffix in {".md", ".markdown"}:
        return "markdown"
    if suffix in TEXT_EXTENSIONS or mime.startswith(TEXT_MIME_PREFIXES) or mime in TEXT_MIME_TYPES:
        return "text"
    raise ArtifactProcessingError(
        "Tipo de arquivo não suportado. Use imagem, .txt, .md, .log, .pdf ou .docx."
    )


def process_artifact(
    *,
    attachment_id: uuid.UUID,
    conversation_id: uuid.UUID,
    filename: str,
    mime_type: str,
    content: bytes,
) -> ArtifactProcessingResult:
    """Extrai texto e devolve chunks pesquisáveis para um artefato."""
    kind = detect_artifact_kind(filename, mime_type)
    if kind == "image":
        return ArtifactProcessingResult(kind=kind, chunks=[])
    if kind == "pdf":
        chunks = _chunks_from_pdf(attachment_id, conversation_id, content)
    elif kind == "docx":
        chunks = _chunks_from_text(
            attachment_id,
            conversation_id,
            _extract_docx(content),
            source_kind=kind,
        )
    elif kind == "log":
        chunks = _chunks_from_log(attachment_id, conversation_id, _decode_text(content))
    else:
        chunks = _chunks_from_text(
            attachment_id,
            conversation_id,
            _decode_text(content),
            source_kind=kind,
        )
    if not chunks:
        raise ArtifactProcessingError("Não foi possível extrair texto pesquisável do arquivo.")
    return ArtifactProcessingResult(kind=kind, chunks=chunks)


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = content.decode(encoding)
            return _normalise_text(text)
        except UnicodeDecodeError:
            continue
    return _normalise_text(content.decode("utf-8", errors="replace"))


def _normalise_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"\r\n?", "\n", text)
    return text.strip()


def _chunks_from_text(
    attachment_id: uuid.UUID,
    conversation_id: uuid.UUID,
    text: str,
    *,
    source_kind: str,
) -> list[ConversationArtifactChunk]:
    text = _normalise_text(text)
    chunks: list[ConversationArtifactChunk] = []
    start = 0
    index = 0
    while start < len(text):
        end = min(start + TEXT_CHUNK_CHARS, len(text))
        if end < len(text):
            window = text[start:end]
            cut = max(window.rfind("\n\n"), window.rfind("\n"), window.rfind(". "))
            if cut > TEXT_CHUNK_CHARS // 2:
                end = start + cut + 1
        chunk_text = text[start:end].strip()
        if chunk_text:
            chunks.append(
                ConversationArtifactChunk(
                    id=uuid.uuid4(),
                    attachment_id=attachment_id,
                    conversation_id=conversation_id,
                    chunk_index=index,
                    content=chunk_text,
                    meta={"kind": source_kind},
                )
            )
            index += 1
        if end >= len(text):
            break
        start = max(end - TEXT_CHUNK_OVERLAP, start + 1)
    return chunks


def _chunks_from_log(
    attachment_id: uuid.UUID,
    conversation_id: uuid.UUID,
    text: str,
) -> list[ConversationArtifactChunk]:
    lines = _normalise_text(text).splitlines()
    chunks: list[ConversationArtifactChunk] = []
    index = 0
    start = 0
    while start < len(lines):
        end = min(start + LOG_CHUNK_LINES, len(lines))
        block = "\n".join(lines[start:end])
        while len(block) > LOG_CHUNK_CHARS and end > start + 1:
            end = start + max(1, (end - start) // 2)
            block = "\n".join(lines[start:end])
        if block.strip():
            chunks.append(
                ConversationArtifactChunk(
                    id=uuid.uuid4(),
                    attachment_id=attachment_id,
                    conversation_id=conversation_id,
                    chunk_index=index,
                    content=block.strip(),
                    line_start=start + 1,
                    line_end=end,
                    meta={"kind": "log"},
                )
            )
            index += 1
        start = end
    return chunks


def _chunks_from_pdf(
    attachment_id: uuid.UUID,
    conversation_id: uuid.UUID,
    content: bytes,
) -> list[ConversationArtifactChunk]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ArtifactProcessingError("pypdf não instalado para processar PDF.") from exc

    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as exc:
        raise ArtifactProcessingError("PDF inválido ou ilegível.") from exc

    chunks: list[ConversationArtifactChunk] = []
    for page_index, page in enumerate(reader.pages):
        try:
            text = _normalise_text(page.extract_text() or "")
        except Exception:
            text = ""
        for partial in _chunks_from_text(
            attachment_id,
            conversation_id,
            text,
            source_kind="pdf",
        ):
            partial.chunk_index = len(chunks)
            partial.page_start = page_index + 1
            partial.page_end = page_index + 1
            chunks.append(partial)
    return chunks


def _extract_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ArtifactProcessingError("python-docx não instalado para processar DOCX.") from exc
    try:
        doc = Document(io.BytesIO(content))
    except Exception as exc:
        raise ArtifactProcessingError("DOCX inválido ou ilegível.") from exc
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _normalise_text("\n".join(parts))
